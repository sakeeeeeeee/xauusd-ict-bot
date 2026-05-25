"""
Generate Word (.docx) and PDF - VERSI LENGKAP panduan upgrade + semua prompt.
Run: python scripts/generate_upgrade_docs.py
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fpdf import FPDF

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "docs"))

from upgrade_full_content import build_sections, build_prompt_sections  # noqa: E402

DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

DOCX_PATH = DOCS / "PANDUAN_UPGRADE_XAUUSD_ICT_BOT_LENGKAP.docx"
PDF_PATH = DOCS / "PANDUAN_UPGRADE_XAUUSD_ICT_BOT_LENGKAP.pdf"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=min(level, 3))


def add_para(doc, text, mono=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if mono:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def render_sections_to_doc(doc, sections, page_break_before_lampiran=True):
    lampiran_started = False
    for title_text, level, lines in sections:
        if page_break_before_lampiran and title_text.startswith("LAMPIRAN") and not lampiran_started:
            doc.add_page_break()
            lampiran_started = True
        add_heading(doc, title_text, level=level)
        for line in lines:
            if isinstance(line, tuple):
                kind, content = line[0], line[1]
                if kind == "bullet":
                    add_bullet(doc, content)
                elif kind.startswith("prompt_"):
                    add_para(doc, content, mono=True)
                    doc.add_paragraph()
            else:
                add_para(doc, line)


def write_docx(all_sections):
    doc = Document()
    t = doc.add_heading("Panduan Upgrade Lengkap", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("XAUUSD ICT Signal Bot")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Versi LENGKAP: saran upgrade + penjelasan + prompt siap copas (A1-L9)")
    doc.add_paragraph("Total prompt item: 70+ | Termasuk batch Fase 1-4")
    doc.add_page_break()

    render_sections_to_doc(doc, all_sections)
    doc.save(DOCX_PATH)
    print(f"DOCX: {DOCX_PATH}")


def write_pdf(all_sections):
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=18)

    def safe(text: str) -> str:
        return (
            text.replace("\u2013", "-")
            .replace("\u2014", "-")
            .replace("\u2192", "->")
            .replace("\u201c", '"')
            .replace("\u201d", '"')
            .encode("latin-1", "replace")
            .decode("latin-1")
        )

    def write_line(text: str, bold: bool = False, size: int = 9):
        pdf.set_font("Helvetica", "B" if bold else "", size)
        for paragraph in text.split("\n"):
            if paragraph.strip():
                pdf.multi_cell(pdf.epw, 4.5, safe(paragraph))
            else:
                pdf.ln(2)

    pdf.add_page()
    write_line("Panduan Upgrade LENGKAP", bold=True, size=16)
    write_line("XAUUSD ICT Signal Bot", size=12)
    write_line("Saran upgrade + prompt copas A1-L9 + batch Fase 1-4")

    lampiran_break = False
    for title_text, level, lines in all_sections:
        if title_text.startswith("LAMPIRAN") and not lampiran_break:
            pdf.add_page()
            lampiran_break = True
        pdf.ln(2)
        write_line(title_text, bold=True, size=13 if level == 1 else 11)
        for line in lines:
            if isinstance(line, tuple):
                kind, content = line[0], line[1]
                if kind == "bullet":
                    write_line("  - " + content, size=8)
                else:
                    write_line(content, size=7)
                    pdf.ln(1)
            else:
                write_line(line, size=8)

    pdf.output(PDF_PATH)
    print(f"PDF:  {PDF_PATH}")


def main():
    all_sections = build_sections() + build_prompt_sections()
    write_docx(all_sections)
    write_pdf(all_sections)
    print("Selesai - VERSI LENGKAP di folder docs/")


if __name__ == "__main__":
    main()
